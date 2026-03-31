"""
number_mask.py — скрыть числа в .docx, потом вернуть обратно

Использование:
  python number_mask.py hide output_2.docx masked.docx map.json
  python number_mask.py restore masked_corrected.docx map.json output.docx
"""

import sys
import re
import json
from docx import Document

# Паттерн: целые и дробные числа (включая разделители: 1 000, 1,000, 1.000)
NUMBER_PATTERN = re.compile(
    r'\b\d[\d\s]*(?:[.,]\d+)*\b'
)

PLACEHOLDER = "{{NUM_{}}}"


def hide_numbers(input_path: str, output_path: str, map_path: str):
    doc = Document(input_path)
    number_map = {}
    counter = [0]

    def replace_in_text(text: str) -> str:
        def replacer(m):
            original = m.group(0)
            key = PLACEHOLDER.format(counter[0])
            number_map[key] = original
            counter[0] += 1
            return key
        return NUMBER_PATTERN.sub(replacer, text)

    def process_paragraph(para):
        for run in para.runs:
            if run.text:
                run.text = replace_in_text(run.text)

    for para in doc.paragraphs:
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    doc.save(output_path)

    with open(map_path, "w", encoding="utf-8") as f:
        json.dump(number_map, f, ensure_ascii=False, indent=2)

    print(f"✅ Скрыто {len(number_map)} чисел")
    print(f"   Маскированный файл: {output_path}")
    print(f"   Карта чисел: {map_path}")


def restore_numbers(masked_path: str, map_path: str, output_path: str):
    doc = Document(masked_path)

    with open(map_path, "r", encoding="utf-8") as f:
        number_map = json.load(f)

    sorted_keys = sorted(number_map.keys(), key=len, reverse=True)

    def restore_in_text(text: str) -> str:
        for key in sorted_keys:
            text = text.replace(key, number_map[key])
        return text

    def process_paragraph(para):
        for run in para.runs:
            if run.text:
                run.text = restore_in_text(run.text)

    for para in doc.paragraphs:
        process_paragraph(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    doc.save(output_path)
    print(f"✅ Числа восстановлены → {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    mode = sys.argv[1]

    if mode == "hide" and len(sys.argv) == 5:
        hide_numbers(sys.argv[2], sys.argv[3], sys.argv[4])
    elif mode == "restore" and len(sys.argv) == 5:
        restore_numbers(sys.argv[2], sys.argv[3], sys.argv[4])
    else:
        print(__doc__)
        sys.exit(1)
